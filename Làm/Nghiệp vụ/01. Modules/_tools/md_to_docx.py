#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Chuyển SRS Markdown → Word (.docx) chuyên nghiệp."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_runs_with_markdown(paragraph, text: str, base_size=11):
    """Hỗ trợ **bold**, `code`, *italic* đơn giản trong một dòng."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=base_size)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 1, color=RGBColor(0x1F, 0x4E, 0x79))
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", s))


def parse_table_row(line: str) -> list[str]:
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf: list[str] = []
    skip_details = False

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip("\n")
        stripped = raw.strip()

        # details/summary — bỏ tag, giữ nội dung
        if stripped.startswith("<details") or stripped.startswith("</details>"):
            skip_details = stripped.startswith("<details")
            i += 1
            continue
        if stripped.startswith("<summary>") and stripped.endswith("</summary>"):
            title = re.sub(r"</?summary>", "", stripped)
            p = doc.add_paragraph()
            add_runs_with_markdown(p, title, base_size=11)
            for run in p.runs:
                run.bold = True
            i += 1
            continue
        if stripped.startswith("<summary>"):
            i += 1
            continue

        # fenced code
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, name="Consolas", size=9, color=RGBColor(0x33, 0x33, 0x33))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                content = m.group(2).strip()
                # Word heading level 1..9; map # -> Title-ish, ## -> H1...
                if level == 1:
                    p = doc.add_heading(content, level=0)
                else:
                    p = doc.add_heading(content, level=min(level - 1, 4))
                for run in p.runs:
                    set_run_font(run, size=16 if level == 1 else (14 if level == 2 else 12), bold=True,
                                 color=RGBColor(0x1F, 0x4E, 0x79))
                i += 1
                continue

        # Blockquote
        if stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            add_runs_with_markdown(p, content)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and not is_table_separator(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(re.sub(r"[*`]", "", h))
                set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.makeelement(
                    qn("w:shd"),
                    {qn("w:fill"): "1F4E79", qn("w:val"): "clear"},
                )
                tcPr.append(shd)
            for r, row in enumerate(rows):
                for c in range(len(headers)):
                    val = row[c] if c < len(row) else ""
                    cell = table.rows[r + 1].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs_with_markdown(p, val, base_size=10)
            doc.add_paragraph()
            continue

        # Lists
        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^([-*]|\d+\.)\s+", "", stripped)
            style_name = "List Number" if re.match(r"^\d+\.", stripped) else "List Bullet"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("• " if style_name == "List Bullet" else "")
            # clear default
            if p.runs:
                p.runs[0].text = ""
            add_runs_with_markdown(p, content)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs_with_markdown(p, stripped)
        i += 1

    # Cover metadata polish: first heading center if title
    if doc.paragraphs:
        first = doc.paragraphs[0]
        if first.style and first.style.name.startswith("Heading"):
            first.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"OK: {docx_path} ({docx_path.stat().st_size // 1024} KB)")


def main():
    if len(sys.argv) < 2:
        print("Usage: md_to_docx.py <file.md> [out.docx]")
        sys.exit(1)
    md = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    md_to_docx(md, out)


if __name__ == "__main__":
    main()
