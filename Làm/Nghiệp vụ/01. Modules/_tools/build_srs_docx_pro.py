#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Xuất SRS Word chuyên nghiệp:
- Trang bìa
- Trang thông tin / kiểm soát tài liệu
- Mục lục (TOC field — cập nhật khi mở Word)
- Nội dung theo heading chuẩn
- Header / Footer + đánh số trang
- Trang kết (phê duyệt / kết thúc tài liệu)
"""
from __future__ import annotations

import argparse
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0B, 0x3D, 0x5C)
MAROON = RGBColor(0x80, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = "D6EAF8"
HEADER_FILL = "0B3D5C"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_md_runs(paragraph, text: str, size=12):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            set_run_font(r, size=size)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run_font(r, size=size, bold=True)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, name="Consolas", size=size - 1, color=NAVY)
        else:
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, size=size, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, size=size)


def shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border_none(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f"<w:tblBorders {nsdecls('w')}>"
        "<w:top w:val='none'/>"
        "<w:left w:val='none'/>"
        "<w:bottom w:val='none'/>"
        "<w:right w:val='none'/>"
        "<w:insideH w:val='none'/>"
        "<w:insideV w:val='none'/>"
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B3D5C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_number_field(paragraph, prefix="Trang "):
    run = paragraph.add_run(prefix)
    set_run_font(run, size=9, color=GRAY)

    def _fld(instr: str):
        r1 = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r1._r.append(fld_begin)

        r2 = paragraph.add_run()
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        r2._r.append(instr_el)

        r3 = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r3._r.append(fld_sep)

        r4 = paragraph.add_run("1")
        set_run_font(r4, size=9, color=GRAY)

        r5 = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r5._r.append(fld_end)

    _fld(" PAGE ")
    run2 = paragraph.add_run(" / ")
    set_run_font(run2, size=9, color=GRAY)
    _fld(" NUMPAGES ")


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run2._r.append(instr)

    run3 = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = p.add_run(
        "[Mục lục sẽ hiển thị tại đây. Trong Microsoft Word: chuột phải → Update Field / Cập nhật trường "
        "→ Update entire table.]"
    )
    set_run_font(run4, size=11, italic=True, color=GRAY)

    run5 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)

    # Force Word to update fields on open
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    for level, size in sizes.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.bold = True
        hs.font.size = Pt(size)
        hs.font.color.rgb = NAVY
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        hs.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        hs.paragraph_format.space_after = Pt(8)


def configure_header_footer(section, doc_code: str, module_name: str, is_cover=False):
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing
    for p in header.paragraphs:
        p.text = ""
    for p in footer.paragraphs:
        p.text = ""

    if is_cover:
        return

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f"{doc_code}  |  {module_name}")
    set_run_font(r, size=9, color=GRAY, italic=True)

    # line under header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "0B3D5C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # top border
    pPrf = fp._p.get_or_add_pPr()
    pBdrf = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "0B3D5C")
    pBdrf.append(top)
    pPrf.append(pBdrf)
    add_page_number_field(fp, prefix="Trang ")
    r2 = fp.add_run("  ·  Bảo mật: Nội bộ dự án")
    set_run_font(r2, size=9, color=GRAY)


def add_cover(doc: Document, meta: dict):
    # Brand line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(meta.get("org", "ERP PRODUCT LINE"))
    set_run_font(r, size=12, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("doc_family_en", "SOFTWARE REQUIREMENTS SPECIFICATION"))
    set_run_font(r, size=11, color=GRAY)

    add_horizontal_line(doc)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("doc_family_vi", "TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM"))
    set_run_font(r, size=14, bold=True, color=MAROON)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(meta["title"])
    set_run_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["subtitle"])
    set_run_font(r, size=14, italic=True, color=GRAY)

    for _ in range(2):
        doc.add_paragraph()

    # Info box
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    data = [
        ("Mã tài liệu", meta["doc_code"]),
        (meta.get("module_label", "Mã module"), meta["module_code"]),
        ("Phiên bản", meta["version"]),
        ("Ngày ban hành", meta["date"]),
        ("Trạng thái", meta["status"]),
        ("Phân loại", meta.get("classification", "Nội bộ — Nghiệp vụ / BA")),
    ]
    for i, (k, v) in enumerate(data):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(5)
        c1.width = Cm(10)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        set_run_font(r0, size=11, bold=True, color=WHITE)
        shade_cell(c0, HEADER_FILL)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        set_run_font(r1, size=11)
        shade_cell(c1, LIGHT)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("footer_note", "Tài liệu dùng để chốt nghiệp vụ trước khi thiết kế & lập trình."))
    set_run_font(r, size=10, italic=True, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("place_date", f"Việt Nam, {meta['date']}"))
    set_run_font(r, size=11, color=GRAY)


def add_doc_control(doc: Document, meta: dict):
    doc.add_heading("1. Thông tin kiểm soát tài liệu", level=1)

    p = doc.add_paragraph()
    add_md_runs(
        p,
        "Trang này ghi nhận thông tin quản lý phiên bản, phân phối và trách nhiệm phê duyệt tài liệu SRS.",
    )

    doc.add_heading("1.1. Thông tin chung", level=2)
    rows = [
        ("Tên tài liệu", meta["title"]),
        ("Mã tài liệu", meta["doc_code"]),
        ("Module", f"{meta['module_code']} — {meta['module_name']}"),
        ("Phiên bản", meta["version"]),
        ("Ngày lập / cập nhật", meta["date"]),
        ("Ngôn ngữ", "Tiếng Việt"),
        ("Định dạng bàn giao", "Microsoft Word (.docx)"),
        ("Trạng thái", meta["status"]),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = ""
        t.rows[i].cells[1].text = ""
        p0 = t.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(k)
        set_run_font(r0, size=11, bold=True)
        shade_cell(t.rows[i].cells[0], LIGHT)
        p1 = t.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(v)
        set_run_font(r1, size=11)

    doc.add_heading("1.2. Lịch sử thay đổi", level=2)
    hist = meta.get("history") or [["1.0", meta["date"], "BA / Solution", "Khởi tạo", meta["status"]]]
    ht = doc.add_table(rows=1 + len(hist), cols=5)
    ht.style = "Table Grid"
    headers = ["Phiên bản", "Ngày", "Người thực hiện", "Mô tả thay đổi", "Trạng thái"]
    for c, h in enumerate(headers):
        cell = ht.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(cell, HEADER_FILL)
    for i, row in enumerate(hist):
        for c, val in enumerate(row):
            cell = ht.rows[i + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=10)

    doc.add_heading("1.3. Danh sách phân phối", level=2)
    dist = meta.get("distribution") or [
        ("Chủ sản phẩm / PO", "Phê duyệt phạm vi"),
        ("Business Analyst", "Biên soạn / cập nhật"),
        ("Solution Architect", "Rà soát khả thi"),
        ("Tech Lead / QA Lead", "Ước lượng & kế hoạch kiểm thử"),
        ("Presales / Triển khai", "Tham chiếu đóng gói bán module"),
    ]
    dt = doc.add_table(rows=1 + len(dist), cols=2)
    dt.style = "Table Grid"
    for c, h in enumerate(["Đối tượng", "Mục đích sử dụng"]):
        cell = dt.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(cell, HEADER_FILL)
    for i, (a, b) in enumerate(dist):
        dt.rows[i + 1].cells[0].text = ""
        dt.rows[i + 1].cells[1].text = ""
        p0 = dt.rows[i + 1].cells[0].paragraphs[0]
        r0 = p0.add_run(a)
        set_run_font(r0, size=10)
        p1 = dt.rows[i + 1].cells[1].paragraphs[0]
        r1 = p1.add_run(b)
        set_run_font(r1, size=10)

    doc.add_page_break()


def add_toc_page(doc: Document):
    doc.add_heading("2. Mục lục", level=1)
    p = doc.add_paragraph()
    add_md_runs(
        p,
        "Mục lục được tạo tự động từ các tiêu đề trong tài liệu. "
        "Khi mở bằng Microsoft Word, hãy **cập nhật trường mục lục** (Right-click → Update Field) nếu Word chưa tự cập nhật.",
        size=11,
    )
    doc.add_paragraph()
    add_toc_field(doc)
    note = doc.add_paragraph()
    r = note.add_run(
        "Gợi ý: View → Outline để kiểm tra cấu trúc Heading 1 / Heading 2 / Heading 3 trước khi phát hành."
    )
    set_run_font(r, size=10, italic=True, color=GRAY)
    doc.add_page_break()


def is_table_separator(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", s))


def parse_table_row(line: str) -> list[str]:
    s = line.strip().strip("|")
    # giữ \| trong ô (escape markdown)
    parts = re.split(r"(?<!\\)\|", s)
    return [c.strip().replace("\\|", "|") for c in parts]


def fill_cell_multiline(cell, text: str, size=10, bold=False, shade=None):
    """Đổ nội dung ô; hỗ trợ <br> / xuống dòng thành nhiều paragraph."""
    cell.text = ""
    if shade:
        shade_cell(cell, shade)
    chunks = re.split(r"<br\s*/?>|\n", text or "")
    if not chunks:
        chunks = [""]
    for idx, chunk in enumerate(chunks):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        # bỏ ** quanh nhãn trường
        clean = chunk.strip()
        is_label_bold = bold or bool(re.match(r"^\*\*.+\*\*$", clean))
        add_md_runs(p, clean, size=size)
        if is_label_bold:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = NAVY


def is_uc_spec_table(headers: list[str], rows: list[list[str]]) -> bool:
    h = " ".join(headers).lower()
    if "trường thông tin" in h and "nội dung" in h:
        return True
    if rows and rows[0] and "use case id" in rows[0][0].lower():
        return True
    return False


def remap_section_number(content: str, offset: int = 2) -> str:
    """Đổi số chương đầu: '1. A' -> '3. A', '1.1. B' -> '3.1. B', '7.2.1. C' -> '9.2.1. C'."""
    mm = re.match(r"^(\d+)((?:\.\d+)*)\.\s+(.*)$", content)
    if not mm:
        return content
    first = int(mm.group(1)) + offset
    return f"{first}{mm.group(2)}. {mm.group(3)}"


def add_body_from_md(doc: Document, md_path: Path, skip_until_heading: str | None = "## 1."):
    """
    Đưa nội dung MD vào Word.
    Bỏ phần front-matter trước heading nội dung chính (mặc định từ ## 1. trở đi của bản v1.1,
    nhưng sẽ remap số chương để khớp cấu trúc chuyên nghiệp: phần nội dung bắt đầu từ chương 3).
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()

    # Tìm điểm bắt đầu nội dung chuyên môn: sau catalog/control của md gốc.
    # Với SRS_SYS_v1.1: bắt đầu từ "## 1. Giới thiệu" nhưng ta sẽ đổi thành chương 3+
    start = 0
    for idx, line in enumerate(lines):
        if line.strip().startswith("## 1. Giới thiệu") or line.strip().startswith("## 1. "):
            start = idx
            break
    else:
        # fallback: sau catalog details đóng
        for idx, line in enumerate(lines):
            if line.strip() == "## 7. Đặc tả Use Case theo nhóm" or line.strip().startswith("## 1."):
                start = idx
                break

    # Nếu file đã là bản đầy đủ, lấy từ "## 1."
    body_lines = lines[start:]

    # Professional structure:
    # 1 Doc control (already added)
    # 2 TOC (already added)
    # 3+ content from original ## 1, ## 2, ... become ## 3, ## 4...
    # Subsections ### 1.1 / #### 1.1.1 cũng +2 ở số chương đầu.

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(body_lines):
        raw = body_lines[i]
        stripped = raw.strip()

        if stripped.startswith("<details") or stripped == "</details>":
            i += 1
            continue
        if stripped.startswith("<summary>"):
            title = re.sub(r"</?summary>", "", stripped)
            p = doc.add_paragraph()
            r = p.add_run(title)
            set_run_font(r, size=11, bold=True, color=NAVY)
            i += 1
            continue

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                set_run_font(r, name="Consolas", size=9)
                p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        # Headings: remap số chương cho ## / ### / #### dạng n. / n.n. / n.n.n.
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            # Skip duplicate doc title (# ...)
            if level == 1 and content.startswith("SRS-"):
                i += 1
                continue

            if level == 2:
                content = remap_section_number(content, 2)
                hlevel = 1
            elif level == 3:
                content = remap_section_number(content, 2)
                hlevel = 2
            elif level == 4:
                content = remap_section_number(content, 2)
                hlevel = 3
            else:
                hlevel = min(level, 4)

            doc.add_heading(content, level=hlevel)
            i += 1
            continue

        if stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_md_runs(p, content, size=11)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = GRAY
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(body_lines) and is_table_separator(body_lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(body_lines) and body_lines[i].strip().startswith("|") and not is_table_separator(body_lines[i]):
                rows.append(parse_table_row(body_lines[i]))
                i += 1

            uc_table = is_uc_spec_table(headers, rows)
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            table.autofit = True

            if uc_table:
                # Header nhẹ + cột trái tô xanh nhạt (đúng mẫu đặc tả UC)
                for c, h in enumerate(headers):
                    fill_cell_multiline(
                        table.rows[0].cells[c],
                        re.sub(r"[*`]", "", h),
                        size=10,
                        bold=True,
                        shade=HEADER_FILL if c == 0 else LIGHT,
                    )
                    if c == 0:
                        for p in table.rows[0].cells[c].paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = WHITE
                for r_i, row in enumerate(rows):
                    for c in range(len(headers)):
                        val = row[c] if c < len(row) else ""
                        fill_cell_multiline(
                            table.rows[r_i + 1].cells[c],
                            val,
                            size=10,
                            bold=(c == 0),
                            shade=LIGHT if c == 0 else None,
                        )
                # Độ rộng gợi ý: nhãn hẹp / nội dung rộng
                if len(headers) == 2:
                    for row in table.rows:
                        row.cells[0].width = Cm(4.2)
                        row.cells[1].width = Cm(12.3)
            else:
                for c, h in enumerate(headers):
                    cell = table.rows[0].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    r = p.add_run(re.sub(r"[*`]", "", h))
                    set_run_font(r, size=10, bold=True, color=WHITE)
                    shade_cell(cell, HEADER_FILL)
                for r_i, row in enumerate(rows):
                    for c in range(len(headers)):
                        val = row[c] if c < len(row) else ""
                        fill_cell_multiline(table.rows[r_i + 1].cells[c], val, size=10)
            doc.add_paragraph()
            continue

        # Caption bảng UC: **Bảng N. Đặc tả Use Case "..."**
        if re.match(r"^\*\*Bảng\s+\d+\.\s+Đặc tả Use Case", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            add_md_runs(p, stripped, size=12)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = NAVY
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^([-*]|\d+\.)\s+", "", stripped)
            try:
                style = "List Number" if re.match(r"^\d+\.", stripped) else "List Bullet"
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
            if p.runs:
                p.runs[0].text = ""
            add_md_runs(p, content, size=12)
            i += 1
            continue

        p = doc.add_paragraph()
        add_md_runs(p, stripped, size=12)
        i += 1


def add_closing_page(doc: Document, meta: dict):
    doc.add_page_break()
    doc.add_heading("Trang kết — Phê duyệt & Đóng tài liệu", level=1)

    p = doc.add_paragraph()
    add_md_runs(
        p,
        "Khi các bên liên quan thống nhất nội dung SRS này, vui lòng ký xác nhận bên dưới. "
        "Sau khi phê duyệt, phiên bản được **đóng băng** để làm căn cứ thiết kế cấu trúc source và phát triển.",
    )

    doc.add_heading("A. Kết luận phạm vi", level=2)
    bullets = meta.get(
        "closing_points",
        [
            "Tài liệu đã mô tả đủ phạm vi In/Out, use case, workflow, BR, NFR và tiêu chí nghiệm thu của module.",
            "Các hạng mục còn mở (nếu có) phải được chốt hoặc chuyển backlog có mã theo dõi trước khi code.",
            "Mọi thay đổi sau phê duyệt phải tạo phiên bản SRS mới và ghi lịch sử thay đổi.",
        ],
    )
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        if p.runs:
            p.runs[0].text = ""
        add_md_runs(p, b)

    doc.add_heading("B. Bảng phê duyệt", level=2)
    roles = meta.get(
        "approvals",
        [
            ("Người biên soạn (BA)", "", "", ""),
            ("Solution Architect", "", "", ""),
            ("Chủ sản phẩm / PO", "", "", ""),
            ("Đại diện kỹ thuật (Tech Lead)", "", "", ""),
        ],
    )
    t = doc.add_table(rows=1 + len(roles), cols=4)
    t.style = "Table Grid"
    for c, h in enumerate(["Vai trò", "Họ và tên", "Chữ ký", "Ngày"]):
        cell = t.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(cell, HEADER_FILL)
    for i, row in enumerate(roles):
        for c in range(4):
            cell = t.rows[i + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            # taller signature rows
            r = p.add_run(row[c] if c < len(row) else "")
            set_run_font(r, size=11)
            if c == 0:
                set_run_font(r, size=11, bold=True)
        # add blank lines in signature cell for writing space
        t.rows[i + 1].cells[2].paragraphs[0].add_run("\n\n\n")

    doc.add_heading("C. Cam kết sau phê duyệt", level=2)
    for b in [
        "Không triển khai trái với SRS đã duyệt trừ khi có Change Request được chấp thuận.",
        "QA lập kịch bản kiểm thử truy vết về mã UC trong tài liệu này.",
        "Bản Word (`.docx`) là bản gốc bàn giao; không dùng bản nháp khác làm căn cứ nghiệm thu.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        if p.runs:
            p.runs[0].text = ""
        add_md_runs(p, b)

    doc.add_paragraph()
    add_horizontal_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— HẾT TÀI LIỆU —")
    set_run_font(r, size=12, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{meta['doc_code']}  ·  Version {meta['version']}  ·  {meta['date']}")
    set_run_font(r, size=10, color=GRAY)


def extract_meta_from_md(md_path: Path) -> dict:
    text = md_path.read_text(encoding="utf-8")
    today = date.today().strftime("%d/%m/%Y")
    m = re.search(r"SRS-([A-Z]+)-v([\d.]+)", text)
    doc_code = m.group(0) if m else "SRS-MOD-v1.1"
    version = m.group(2) if m else "1.1"
    module_code = m.group(1) if m else "MOD"

    title_m = re.search(r"^#\s+(.+)$", text, re.M)
    title_line = title_m.group(1) if title_m else f"{doc_code}"
    if "—" in title_line:
        nice = title_line.split("—", 1)[1].strip()
    else:
        nice = title_line

    return {
        "org": "ERP MODULAR PRODUCT",
        "title": nice,
        "subtitle": f"Software Requirements Specification — Module {module_code}",
        "doc_code": doc_code,
        "module_code": module_code,
        "module_name": nice,
        "version": version,
        "date": today,
        "status": "Chờ duyệt nghiệp vụ",
        "classification": "Nội bộ dự án — BA / Solution",
        "place_date": f"Việt Nam, {today}",
        "footer_note": "Tài liệu chốt nghiệp vụ trước khi thiết kế cấu trúc source & lập trình.",
        "history": [
            ["1.0", today, "BA / Generator", "Khởi tạo hàng loạt từ catalog", "Thay thế"],
            ["1.1", today, "BA / Solution", "Viết lại đặc tả UC chuyên sâu + chuẩn hóa Word", "Chờ duyệt"],
        ],
    }


# backward-compatible alias
extract_meta_from_sys_md = extract_meta_from_md


def build(md_path: Path, out_path: Path, meta: dict | None = None):
    meta = meta or extract_meta_from_md(md_path)
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)

    setup_styles(doc)

    # Section 0: trang bìa — không đánh số trang
    cover_section = doc.sections[0]
    configure_header_footer(cover_section, meta["doc_code"], meta["module_name"], is_cover=True)
    add_cover(doc, meta)

    # Section 1: nội dung — đánh số trang từ 1
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin = Cm(2.2)
    body_section.bottom_margin = Cm(2.2)
    body_section.left_margin = Cm(2.5)
    body_section.right_margin = Cm(2.5)
    sectPr = body_section._sectPr
    # remove existing pgNumType if any then set start=1
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)
    configure_header_footer(body_section, meta["doc_code"], meta["module_name"], is_cover=False)

    add_doc_control(doc, meta)
    add_toc_page(doc)
    add_body_from_md(doc, md_path)
    add_closing_page(doc, meta)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote professional DOCX: {out_path} ({out_path.stat().st_size // 1024} KB)")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md", type=Path)
    ap.add_argument("-o", "--out", type=Path, default=None)
    args = ap.parse_args()
    md = args.md
    out = args.out or md.with_name(md.stem + "_PRO.docx")
    # For SYS default overwrite the main docx name
    if out is None:
        out = md.with_suffix(".docx")
    build(md, out)


if __name__ == "__main__":
    main()
